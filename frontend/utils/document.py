"""
Stateless document formatting and byte buffer generators.
"""

import io

from docx import Document

__all__ = ["build_final_document_string", "create_docx_buffer"]


def build_final_document_string(
    *, extracted_answers: dict[str, str], missing_questions: list[str]
) -> str:
    """
    Aggregates text chunks into a markdown document string.
    """
    document_blocks: list[str] = ["# Final Extracted Document\n\n"]

    for question_text, answer_text in extracted_answers.items():
        document_blocks.append(f"### {question_text}\n{answer_text}\n\n")

    unanswered = [
        question for question in missing_questions if question not in extracted_answers
    ]
    for question_text in unanswered:
        document_blocks.append(f"### {question_text}\n*No answer provided*\n\n")

    return "".join(document_blocks)


def create_docx_buffer(
    *, extracted_answers: dict[str, str], missing_questions: list[str]
) -> bytes:
    """
    Generates a Microsoft Word (.docx) document in memory and returns bytes.
    """
    doc = Document()
    doc.add_heading("Final Extracted Document", level=0)

    for question_text, answer_text in extracted_answers.items():
        doc.add_heading(question_text, level=2)
        doc.add_paragraph(answer_text)

    unanswered = [
        question for question in missing_questions if question not in extracted_answers
    ]
    for question_text in unanswered:
        doc.add_heading(question_text, level=2)
        p = doc.add_paragraph()
        p.add_run("No answer provided").italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
