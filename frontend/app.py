"""
Primary streamlit rendering
"""

import io
import logging
from typing import Final, cast

import streamlit as st
from docx import Document

from frontend.api import fetch_all_historical_tasks, fetch_server_templates
from frontend.components.results import (
    render_answers_and_missing_sections,
    render_trust_audit_ledger,
)
from frontend.components.sidebar import (
    purge_active_view,
    render_historical_sidebar,
)
from frontend.config import (
    MODEL_CONFIGURATIONS,
    TEMPLATE_DESCRIPTIONS,
    TEMPLATE_DISPLAY_NAMES,
)
from frontend.protocols import UploadedFileProtocol
from frontend.services import send_audit_request, send_generation_request

__all__ = ["main"]

logger: Final[logging.Logger] = logging.getLogger(__name__)
OVERVIEW_PAGE: Final[str] = "OVERVIEW"


def _initialize_session_state() -> None:
    """
    Set up core streamlit session with explicit mutation
    """

    defaults: dict[str, bool | dict[str, object] | str | None] = {
        "generator_report": None,
        "source_context": None,
        "audit_metrics": None,
        "job_running": False,
        "current_task_id": None,
        "current_task_custom_name": None,
        "historical_audits": {},
        "run_state": "idle",
        "selected_template": OVERVIEW_PAGE,
    }

    for key, value in defaults.items():
        st.session_state.setdefault(key, value)


def _handle_pending_generation(generation_args: object) -> None:
    if not isinstance(generation_args, dict):
        return
    send_generation_request(
        target_document=cast(str, generation_args.get("target_document", "")),
        chosen_engine=cast(str, generation_args.get("chosen_engine", "")),
        uploaded_files=cast(
            list[UploadedFileProtocol],
            generation_args.get("uploaded_files", []),
        ),
        custom_name=cast(str, generation_args.get("custom_name", "")),
    )
    st.session_state.job_running = False
    st.rerun()


def _handle_pending_audit(audit_args: object) -> None:
    if not isinstance(audit_args, dict):
        return
    args_copy = dict(audit_args)
    task_id: str = str(args_copy.pop("task_id", ""))
    metrics = send_audit_request(
        chosen_engine=cast(str, args_copy.get("chosen_engine", "")),
        answers=cast(dict[str, str], args_copy.get("answers", {})),
        judge_iterations=cast(int, args_copy.get("judge_iterations", 3)),
        source_context=cast(str, args_copy.get("source_context", "")),
    )
    st.session_state.job_running = False
    if not metrics:
        st.rerun()
        return

    st.session_state.audit_metrics = metrics
    historical_audits = st.session_state.get("historical_audits")
    if not isinstance(historical_audits, dict):
        historical_audits = {}
    historical_audits[task_id] = metrics
    st.session_state.historical_audits = historical_audits
    st.rerun()


def _should_execute_pending_job() -> bool:
    if not st.session_state.get("job_running"):
        return False
    run_state = st.session_state.get("run_state", "idle")
    if run_state == "triggered":
        st.session_state.run_state = "executing"
        return False
    if run_state != "executing":
        return False
    st.session_state.run_state = "idle"
    return True


def _process_pending_jobs() -> None:
    """
    Executes queded background tasks
    then refreshes the app
    """

    if not _should_execute_pending_job():
        return

    if "pending_generation" in st.session_state:
        _handle_pending_generation(st.session_state.pop("pending_generation"))
        return

    if "pending_audit" in st.session_state:
        _handle_pending_audit(st.session_state.pop("pending_audit"))
        return


def _purge_workspace_heap() -> None:
    """
    Removes active tracking keys to reset the UI
    to baseline so things don't get messy
    """
    purge_active_view()


def _render_template_button(
    template_name: str, selected_page: str, disabled: bool
) -> None:
    btn_label = TEMPLATE_DISPLAY_NAMES.get(template_name, template_name)
    is_active = template_name == selected_page
    clicked = st.sidebar.button(
        btn_label,
        key=f"template_page_{template_name}",
        type="primary" if is_active else "secondary",
        width="stretch",
        disabled=disabled,
    )
    if not clicked:
        return
    st.session_state.selected_template = template_name
    st.rerun()


def _render_sidebar_navigation(*, disabled: bool, templates: list[str]) -> str:
    """
    Renders the sidebar page navigation: one page per form template.
    Returns the active page (OVERVIEW_PAGE or a template key)
    """
    selected_page: str = st.session_state.get("selected_template", OVERVIEW_PAGE)
    is_invalid = selected_page != OVERVIEW_PAGE and selected_page not in templates
    if is_invalid:
        selected_page = OVERVIEW_PAGE
        st.session_state.selected_template = selected_page

    for template_name in templates:
        _render_template_button(template_name, selected_page, disabled)

    return selected_page


def _render_overview_page() -> None:
    """
    Renders a stunning, premium landing page detailing the features
    and offering a primary contextual action button to get started.
    """

    st.markdown(
        """
        <div style="margin-top: 1rem; margin-bottom: 2rem;">
            <h1 style="font-size: 2.5rem; font-weight: 800;
                background: linear-gradient(95deg, #2563eb, #3b82f6, #1d4ed8);
                -webkit-background-clip: text;
                -webkit-text-fill-color: transparent; margin-bottom: 0.5rem;">
                ESM Data Automation Pipeline
            </h1>
            <p style="font-size: 1.1rem; color: #4b5563; line-height: 1.6;
                max-width: 800px;">
                Accelerate scientific data stewardship. Extract structured
                metadata and build high-quality dataset documentation
                (READMEs, NOAA Data Management Plans) directly from
                publications and source files.
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.markdown("### Key Capabilities")

    col1, col2 = st.columns(2)

    with col1:
        st.markdown(
            """
            <div style="background-color: #f8fafc; border: 1px solid #e2e8f0;
                border-radius: 12px; padding: 20px; margin-bottom: 20px;
                min-height: 180px; transition: transform 0.2s ease;">
                <h4 style="margin-top: 0; color: #1e3a8a; display: flex;
                    align-items: center; gap: 8px;">
                    Document Generator
                </h4>
                <p style="color: #475569; font-size: 0.9rem; line-height: 1.5;
                    margin-bottom: 0;">
                    Upload dataset files, scientific papers, or metadata dumps
                    to generate standard formats. Edit answers on the fly,
                    review missing information, and export files as Word
                    (.docx) or Markdown (.md).
                </p>
            </div>
            """,
            unsafe_allow_html=True,
        )

    with col2:
        st.markdown(
            """
            <div style="background-color: #f8fafc; border: 1px solid #e2e8f0;
                border-radius: 12px; padding: 20px; margin-bottom: 20px;
                min-height: 180px; transition: transform 0.2s ease;">
                <h4 style="margin-top: 0; color: #1e3a8a; display: flex;
                    align-items: center; gap: 8px;">
                    LLM Judge Evaluation
                </h4>
                <p style="color: #475569; font-size: 0.9rem; line-height: 1.5;
                    margin-bottom: 0;">
                    Audit and score historical extraction runs. Run
                    multi-iteration stability tests to compute item-level
                    agreement scores (Gwet's AC1) and verify accuracy against
                    the source documents.
                </p>
            </div>
            """,
            unsafe_allow_html=True,
        )

    st.markdown("### Workflow Overview")
    st.markdown(
        """
        1. **Select a template** from the sidebar navigation
           (e.g. README or Data Management Plan).
        2. **Upload reference files** (scientific text files, publications,
           netCDF headers) and trigger AI generation.
        3. **Review and edit** the generated fields directly in the tabbed
           interface to fill in missing details.
        4. **Download** your final curated document in Markdown or Word.
        5. **Audit the run** in the LLM Judge tab to inspect agreement
           scores and ensure output stability.
        """
    )

    st.markdown(
        "<div style='margin-top: 2rem; margin-bottom: 1rem;'>",
        unsafe_allow_html=True,
    )
    if st.button("Start Generating Documentation", type="primary"):
        st.session_state.selected_template = "README"
        st.rerun()
    st.markdown("</div>", unsafe_allow_html=True)


def _render_step_one_upload(*, disabled: bool, target_document: str) -> str:
    """
    Renders the step 1 upload form aligned with page descriptions.
    """
    display_name = TEMPLATE_DISPLAY_NAMES.get(target_document, target_document)
    st.header(f"1. Generate {display_name}")

    form_description = TEMPLATE_DESCRIPTIONS.get(target_document)
    if form_description:
        st.markdown(form_description)

    uploaded_files = st.file_uploader(
        "Drop your scientific data, READMEs, publications, ect... here:",
        accept_multiple_files=True,
        disabled=disabled,
    )

    custom_name = st.text_input(
        "Label this extraction run (optional):",
        placeholder="Project #1",
        disabled=disabled,
    )

    # Submission button style changes dynamically once file is uploaded
    button_type = "primary" if uploaded_files else "secondary"
    trigger_generation = st.button(
        "Read Files & Write Answers",
        type=button_type,
        use_container_width=True,
        disabled=not uploaded_files or disabled,
    )

    if not trigger_generation:
        return target_document

    st.session_state.job_running = True
    st.session_state.run_state = "triggered"
    st.session_state.pending_generation = {
        "target_document": target_document,
        "chosen_engine": st.session_state.get("global_chosen_engine", "Gemini"),
        "uploaded_files": uploaded_files,
        "custom_name": custom_name,
    }
    st.rerun()
    return None


def _build_final_document_string(
    *, extracted_answers: dict[str, str], missing_questions: list[str]
) -> str:
    """
    Aggregates text chunks
    """

    document_blocks: list[str] = ["# Final Extracted Document\n\n"]

    for question_text, answer_text in extracted_answers.items():
        document_blocks.append(f"### {question_text}\n{answer_text}\n\n")

    unanswered = [q for q in missing_questions if q not in extracted_answers]
    for question_text in unanswered:
        document_blocks.append(f"### {question_text}\n*No answer provided*\n\n")

    return "".join(document_blocks)


def _create_docx_buffer(
    *, extracted_answers: dict[str, str], missing_questions: list[str]
) -> bytes:
    """
    Generates a Microsoft Word (.docx) document in memory and returns bytes.
    """

    doc = Document()
    doc.add_heading("Final Extracted Document", level=0)

    for question_text, answer_text in extracted_answers.items():
        doc.add_heading(question_text, level=2)
        doc.add_paragraph(answer_text)

    unanswered = [q for q in missing_questions if q not in extracted_answers]
    for question_text in unanswered:
        doc.add_heading(question_text, level=2)
        p = doc.add_paragraph()
        p.add_run("No answer provided").italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _render_step_three_download(
    *,
    target_document: str,
    extracted: dict[str, str],
    missing: list[str],
    disabled: bool,
) -> None:
    """
    Provides the final aggregated document for download
    """

    st.header("3. Download Final Document")

    download_format = st.radio(
        "Choose Download Format",
        options=["Markdown (.md)", "Microsoft Word (.docx)"],
        horizontal=True,
        disabled=disabled,
    )

    base_name = f"{target_document}_completed"
    custom_name = st.session_state.get("current_task_custom_name")
    if isinstance(custom_name, str) and custom_name.strip():
        cleaned_name = custom_name.strip()
        for char in r'\/:*?"<>|':
            cleaned_name = cleaned_name.replace(char, "_")
        base_name = cleaned_name

    if download_format == "Markdown (.md)":
        final_markdown: str = _build_final_document_string(
            extracted_answers=extracted, missing_questions=missing
        )
        st.download_button(
            label="Download Document (.md)",
            data=final_markdown,
            file_name=f"{base_name}.md",
            mime="text/markdown",
            type="primary",
            disabled=disabled,
        )
    else:
        final_docx: bytes = _create_docx_buffer(
            extracted_answers=extracted, missing_questions=missing
        )
        st.download_button(
            label="Download Document (.docx)",
            data=final_docx,
            file_name=f"{base_name}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            disabled=disabled,
        )


def _render_generator_tab(*, is_running: bool, target_document: str) -> None:
    """
    Renders document filling process
    """

    _render_step_one_upload(
        disabled=is_running,
        target_document=target_document,
    )

    report = cast(dict[str, object] | None, st.session_state.get("generator_report"))
    if not report:
        return

    render_answers_and_missing_sections(disabled=is_running)
    st.markdown("---")

    missing_question: list[str] = cast(list[str], report.get("missing_information", []))
    extracted_answers: dict[str, str] = cast(
        dict[str, str], report.get("extracted_answers", {})
    )

    _render_step_three_download(
        target_document=target_document,
        extracted=extracted_answers,
        missing=missing_question,
        disabled=is_running,
    )

    audit_metrics = cast(
        dict[str, object] | None, st.session_state.get("audit_metrics")
    )
    if not audit_metrics:
        return

    st.markdown("---")
    st.subheader("Complete run snapshot")
    metadata = cast(dict[str, object], audit_metrics.get("metadata", {}))
    kappa_score = cast(
        float,
        metadata.get("global_gwet_ac1") or metadata.get("global_gwets_ac1", 0.0),
    )

    st.metric("Agreement score (Gwet's AC1)", f"{kappa_score:.3f}")
    st.dataframe(
        cast(list[object], audit_metrics.get("item_level_stability_metrics", [])),
        use_container_width=True,
    )


def _get_currently_active_task(
    historical_tasks: list[dict[str, object]], task_id: str | None
) -> dict[str, object] | None:
    if not task_id:
        return None

    valid_tasks = [
        task
        for task in historical_tasks
        if task.get("status") == "COMPLETED" and task.get("report") is not None
    ]

    matching = [t for t in valid_tasks if str(t.get("task_id")) == str(task_id)]
    if not matching:
        return None

    return matching[0]


def _trigger_stability_test(
    task_id: str,
    active_task_data: dict[str, object],
    chosen_engine: str,
    judge_iterations: int,
) -> None:
    report_data = active_task_data.get("report") or {}

    extracted_answers_dict = {}
    if isinstance(report_data, dict):
        extracted_answers_dict = report_data.get("extracted_answers", {})

    st.session_state.job_running = True
    st.session_state.run_state = "triggered"
    st.session_state.pending_audit = {
        "task_id": task_id,
        "chosen_engine": chosen_engine,
        "judge_iterations": judge_iterations,
        "answers": extracted_answers_dict,
        "source_context": active_task_data.get("source_context", ""),
    }
    st.rerun()


def _render_audit_results(audit_metrics: dict[str, object] | None) -> None:
    if not audit_metrics:
        return
    st.markdown("---")
    st.success("Audit complete!")

    metadata = cast(dict[str, object], audit_metrics.get("metadata", {}))
    kappa_score = cast(
        float,
        metadata.get("global_gwet_ac1") or metadata.get("global_gwets_ac1", 0.0),
    )

    st.metric("Agreement score (Gwet's AC1)", f"{kappa_score:.3f}")
    st.dataframe(
        cast(list[object], audit_metrics.get("item_level_stability_metrics", [])),
        use_container_width=True,
    )


def _render_judge_tab(*, disabled: bool, models: list[str]) -> None:
    """
    Renders the LLM Judge tab using the globally selected active run.
    """

    st.header("LLM Judge: Evaluate Historical Extraction")
    st.markdown(
        "Quantify the extraction accuracy of a past run against its "
        "original source context."
    )

    historical_tasks: list[dict[str, object]] = fetch_all_historical_tasks()
    currently_selected_task_id: str | None = st.session_state.get("current_task_id")
    currently_active_task_data = _get_currently_active_task(
        historical_tasks, currently_selected_task_id
    )

    if not currently_active_task_data:
        st.info(
            "Please upload files under the 'Document Generator' tab or "
            "select a past run from the sidebar to evaluate."
        )
        return

    active_run_custom_name = (
        currently_active_task_data.get("custom_name") or "Unnamed Run"
    )
    st.success(
        "Evaluating Active Run: "
        f"**{active_run_custom_name}** (ID: `{str(currently_selected_task_id)[:8]}`)"
    )

    chosen_engine: str = st.selectbox(
        "Select Evaluating AI Judge", models, disabled=disabled
    )

    judge_iterations: int = st.slider(
        "Testing Iterations (Higher = more accurate but much slower!)",
        min_value=2,
        max_value=10,
        value=3,
        disabled=disabled,
    )

    st.markdown("#### Original Source Documents Under Review")
    render_trust_audit_ledger(
        source_context=cast(
            str | None, currently_active_task_data.get("source_context")
        )
    )

    if st.button("Run Stability Test", type="primary", disabled=disabled):
        _trigger_stability_test(
            cast(str, currently_selected_task_id),
            currently_active_task_data,
            chosen_engine,
            judge_iterations,
        )

    audit_metrics = cast(
        dict[str, object] | None, st.session_state.get("audit_metrics")
    )
    _render_audit_results(audit_metrics)


def _render_active_run_header(currently_active_task_id: str | None) -> None:
    if not currently_active_task_id:
        return

    task_custom_name = st.session_state.get("current_task_custom_name")
    short_task_id = str(currently_active_task_id)[:8]
    custom_name = task_custom_name or f"Job {short_task_id}"
    st.markdown(f"### Active Run: **{custom_name}**")


def _render_sidebar(
    is_running: bool,
    available_templates: list[str],
    available_models: list[str],
) -> str:
    """
    Renders the sidebar navigation and returns the selected page.
    """
    selected_page: str = st.session_state.get("selected_template", OVERVIEW_PAGE)
    is_invalid = (
        selected_page != OVERVIEW_PAGE and selected_page not in available_templates
    )
    if is_invalid:
        selected_page = OVERVIEW_PAGE
        st.session_state.selected_template = selected_page

    # The Overview button is rendered at the very top of the sidebar
    if st.sidebar.button(
        "Overview",
        key="page_overview",
        type="primary" if selected_page == OVERVIEW_PAGE else "secondary",
        width="stretch",
        disabled=is_running,
    ):
        st.session_state.selected_template = OVERVIEW_PAGE
        st.rerun()

    # Form Templates Section Title
    st.sidebar.markdown(
        "<div style='font-size: 0.75rem; font-weight: 700; "
        "letter-spacing: 0.05em; color: #6b7280; text-transform: uppercase; "
        "margin-top: 1.5rem; margin-bottom: 0.5rem; padding-left: 12px;'>"
        "Form Templates</div>",
        unsafe_allow_html=True,
    )

    selected_page = _render_sidebar_navigation(
        disabled=is_running,
        templates=available_templates,
    )

    # Consolidated Session Management is rendered next
    render_historical_sidebar()

    # Settings Section Title
    st.sidebar.markdown(
        "<div style='font-size: 0.75rem; font-weight: 700; "
        "letter-spacing: 0.05em; color: #6b7280; text-transform: uppercase; "
        "margin-top: 2rem; margin-bottom: 0.5rem; padding-left: 12px;'>"
        "Settings</div>",
        unsafe_allow_html=True,
    )
    st.sidebar.selectbox(
        "Select AI Model",
        available_models,
        disabled=is_running,
        key="global_chosen_engine",
    )
    return selected_page


def main() -> None:
    """
    Main control flow
    """

    st.set_page_config(page_title="ESM Data Automation", layout="wide")

    # Inject custom CSS to fix design flaws and enhance UI aesthetics
    st.markdown(
        """
        <style>
        /* 1. Global typography and line height */
        div.stMarkdown p,
        div.stMarkdown li {
            line-height: 1.65 !important;
            font-size: 0.95rem !important;
            color: #374151 !important;
        }

        /* Limit reading text layout width for better ergonomics */
        div.stMarkdown {
            max-width: 800px !important;
        }

        /* Breathing room for headers */
        div.stMarkdown h1,
        div.stMarkdown h2,
        div.stMarkdown h3,
        div.stMarkdown h4 {
            margin-top: 1.75rem !important;
            margin-bottom: 0.85rem !important;
            font-weight: 600 !important;
            color: #111827 !important;
        }

        /* 2. Sidebar Navigation Button Styling (Borderless List Items) */
        .st-key-page_overview button,
        div[class*="st-key-template_page_"] button {
            width: 100% !important;
            border: none !important;
            text-align: left !important;
            justify-content: flex-start !important;
            font-weight: 500 !important;
            transition: all 0.2s ease !important;
            margin: 2px 0 !important;
            height: auto !important;
            padding: 8px 12px !important;
            box-shadow: none !important;
            border-radius: 6px !important;
        }

        /* Inactive Sidebar Navigation Items */
        .st-key-page_overview button[data-testid="baseButton-secondary"],
        div[class*="st-key-template_page_"] button[data-testid="baseButton-secondary"] {
            background-color: transparent !important;
            color: #4b5563 !important;
            border: none !important;
            box-shadow: none !important;
        }
        .st-key-page_overview button[data-testid="baseButton-secondary"]:hover,
        div[class*="st-key-template_page_"]
        button[data-testid="baseButton-secondary"]:hover {
            background-color: #f3f4f6 !important;
            color: #1f2937 !important;
            border: none !important;
            box-shadow: none !important;
        }

        /* Active Sidebar Navigation Items (Crisp Blue Accent & Left Border) */
        .st-key-page_overview button[data-testid="baseButton-primary"],
        div[class*="st-key-template_page_"] button[data-testid="baseButton-primary"] {
            background-color: #eff6ff !important;
            color: #2563eb !important;
            border: none !important;
            border-left: 4px solid #2563eb !important;
            border-radius: 0 6px 6px 0 !important;
            padding-left: 8px !important;
            font-weight: 600 !important;
            box-shadow: none !important;
        }

        /* 3. Session Management Sidebar Buttons (+ New / Delete Current) */
        .st-key-new_run_sidebar_btn button,
        .st-key-delete_run_sidebar_btn button {
            background-color: transparent !important;
            color: #6b7280 !important;
            border: none !important;
            border-radius: 6px !important;
            font-weight: 500 !important;
            transition: all 0.2s ease !important;
            box-shadow: none !important;
        }
        .st-key-new_run_sidebar_btn button:hover {
            background-color: #f3f4f6 !important;
            color: #1f2937 !important;
        }
        .st-key-delete_run_sidebar_btn button:hover {
            background-color: #fee2e2 !important;
            color: #dc2626 !important;
        }
        .st-key-new_run_sidebar_btn button:disabled,
        .st-key-delete_run_sidebar_btn button:disabled {
            color: #d1d5db !important;
            background-color: transparent !important;
            cursor: not-allowed !important;
        }

        /* 4. Form Width Alignment Constraints */
        div[data-testid="stFileUploader"],
        div[data-testid="stTextInput"],
        div[data-testid="stForm"],
        div[data-testid="stMain"] div[data-testid="stSelectbox"],
        div[data-testid="stMain"] div[data-testid="stSlider"],
        div[data-testid="stMain"] div[data-testid="stRadio"],
        div[data-testid="stMain"] div.stButton button,
        div[data-testid="stMain"] div.stDownloadButton button {
            max-width: 800px !important;
        }

        /* 4. Tab Component Styling (Pill-shaped iOS style Container, Capped Width) */
        div[data-baseweb="tab-list"] {
            background-color: #f3f4f6 !important;
            padding: 4px !important;
            border-radius: 8px !important;
            border-bottom: none !important;
            gap: 4px !important;
            display: inline-flex !important;
            width: 100% !important;
            max-width: 800px !important;
        }
        div[data-baseweb="tab-list"] button {
            background-color: transparent !important;
            border: none !important;
            border-radius: 6px !important;
            color: #4b5563 !important;
            padding: 8px 16px !important;
            font-weight: 500 !important;
            transition: all 0.2s ease !important;
            flex: 1 1 0% !important;
            text-align: center !important;
        }
        div[data-baseweb="tab-list"] button:hover {
            color: #1f2937 !important;
            background-color: rgba(255, 255, 255, 0.5) !important;
        }
        div[data-baseweb="tab-list"] button[aria-selected="true"] {
            background-color: white !important;
            color: #2563eb !important;
            box-shadow: 0 1px 3px rgba(0, 0, 0, 0.1) !important;
            font-weight: 600 !important;
            border-bottom: none !important;
        }
        div[data-baseweb="tab-highlight"] {
            display: none !important;
        }

        /* 5. Streamlit Default Deploy Button Styling (Solid Blue Button) */
        [data-testid="stHeaderDeployButton"] button,
        [data-testid="stHeaderDeployButton"] a {
            background-color: #2563eb !important;
            color: white !important;
            border: 1px solid #2563eb !important;
            border-radius: 6px !important;
            padding: 6px 14px !important;
            font-weight: 600 !important;
            font-size: 0.85rem !important;
            text-decoration: none !important;
            display: inline-flex !important;
            align-items: center !important;
            justify-content: center !important;
            transition: all 0.2s ease !important;
            box-shadow: 0 1px 2px rgba(0, 0, 0, 0.05) !important;
            height: 32px !important;
            margin-right: 8px !important;
        }
        [data-testid="stHeaderDeployButton"] button:hover,
        [data-testid="stHeaderDeployButton"] a:hover {
            background-color: #1d4ed8 !important;
            border-color: #1d4ed8 !important;
            color: white !important;
            box-shadow: 0 4px 6px -1px rgba(37, 99, 235, 0.2),
                0 2px 4px -1px rgba(37, 99, 235, 0.1) !important;
        }

        /* 6. Sidebar Alignment and Spacing (Perfect Left-Alignment) */
        div[data-testid="stSidebar"] h2,
        div[data-testid="stSidebar"] h3,
        div[data-testid="stSidebar"] h4,
        div[data-testid="stSidebar"] [data-testid="stMarkdownContainer"] p,
        div[data-testid="stSidebar"] .stCaptionContainer {
            padding-left: 12px !important;
            padding-right: 12px !important;
            margin-left: 0 !important;
        }
        div[data-testid="stSidebar"] div[data-testid="stButton"],
        div[data-testid="stSidebar"] div[data-testid="stSelectbox"] {
            padding-left: 12px !important;
            padding-right: 12px !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )
    st.title("ESM Data Automation Pipeline")

    _initialize_session_state()
    _process_pending_jobs()

    is_running: bool = bool(st.session_state.get("job_running"))

    if is_running:
        st.warning("Active AI job currently running...")

    available_templates: list[str] = fetch_server_templates()
    available_models: list[str] = list(MODEL_CONFIGURATIONS.keys())
    selected_page: str = _render_sidebar(
        is_running, available_templates, available_models
    )

    if selected_page == OVERVIEW_PAGE:
        _render_overview_page()
        return

    # Main Workspace header area: show active run name and delete popover if loaded
    currently_active_task_id = st.session_state.get("current_task_id")
    _render_active_run_header(currently_active_task_id)

    tab_generator, tab_judge = st.tabs(["Document Generator", "LLM Judge Evaluation"])

    with tab_generator:
        _render_generator_tab(
            is_running=is_running,
            target_document=selected_page,
        )

    with tab_judge:
        _render_judge_tab(
            disabled=is_running,
            models=available_models,
        )

    if st.session_state.get("run_state") == "executing":
        st.rerun()


if __name__ == "__main__":
    main()
